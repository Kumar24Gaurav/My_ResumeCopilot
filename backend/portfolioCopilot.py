"""
# project setup import packages and setting llm client and model

# pdf read -> raw text

# pdf parse

# metadata (links)
# Tools
      ├── search_resume()
      └── get_link()

# tool definitions
# chat_with_resume()

# agent
# finish
"""

import os
from pathlib import Path
from groq import Groq
from pypdf import PdfReader
from docx import Document
from dotenv import load_dotenv
import json

load_dotenv()

my_api_key = os.getenv("GROQ_API_KEY")
if not my_api_key:
    raise ValueError ("api error")

client = Groq(api_key=my_api_key)
model = "openai/gpt-oss-120b"

resume_path = Path("resume")/"kumargaurav.pdf"
if resume_path.suffix.lower() not in [".pdf",".docx"]:
    raise ValueError ("only pdf or docx files are allowed")

def read_resume(resume_path):
    if resume_path.suffix.lower() == ".pdf":
        return read_pdf(resume_path)
    elif resume_path.suffix.lower() == ".docx":
        return read_docx(resume_path)
    else:
        return None

def read_pdf(resume_path):
    reader = PdfReader(resume_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(resume_path):
    document = Document(resume_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


resume_text = read_resume(resume_path)

from pydantic import BaseModel
from typing import List, Optional
class Contact(BaseModel):
    email: str
    phone: str

class Education(BaseModel):
    degree: str
    college: str
    score: Optional[str] = None

class Experience(BaseModel):
    company: str
    role: str
    duration: Optional[str] = None
    description: Optional[str] = None

class Project(BaseModel):
    name: str
    description: str
    tech_stack: List[str]

class Resume(BaseModel):
    name: str
    summary: Optional[str] = None
    contact: Contact
    skills: List[str]
    education: List[Education]
    experience: List[Experience]
    projects: List[Project]
    certifications: Optional[List[str]] = None


resume_schema = Resume.model_json_schema()

def extract_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships
    - certification

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Return only JSON.

    Do not wrap the JSON inside markdown.

    Do not explain anything.

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume


resume = extract_resume(resume_text)

# print(resume.model_dump_json(indent=4))


METADATA = {

    "links": {
        "portfolio": "https://kumargaurav-portfolio.vercel.app",

        "github": "https://github.com/Kumar24Gaurav",

        "linkedin": "https://www.linkedin.com/in/kumar-gaurav-814a58299"
    },

    "projects":{

        "finsight":{

            "live":"https://fin-sight-sandy.vercel.app",

            "github":"https://github.com/Kumar24Gaurav/FinSight"
        },

        "resume evaluator":{

            "github":"https://github.com/Kumar24Gaurav/AI_Resume_Evaluator"
        }

    }

}


def get_link(query: str):

    query = query.lower().strip()

    links = METADATA["links"]

    for key, value in links.items():
        if key in query:
            return value

    projects = METADATA["projects"]

    for project_name in projects:
        if project_name in query:
            return projects[project_name]

    if "project" in query:
        return projects

    return None


def search_resume(query: str):

    query = query.lower()

    if any(word in query for word in ["skill", "technology", "tech stack"]):
        return resume.skills

    if "education" in query:
        return resume.education

    if "experience" in query:
        return resume.experience

    if any(word in query for word in ["certification", "certificate", "certifications"]):
        return resume.certifications

    if "project" in query:
        return resume.projects

    if "email" in query:
        return resume.contact.email

    if "phone" in query:
        return resume.contact.phone

    if "contact" in query:
        return resume.contact

    if any(word in query for word in ["summary", "about", "profile", "introduce"]):
        if resume.summary:
            return resume.summary

        return resume.model_dump()

    for project in resume.projects:
        if query in project.name.lower():
            return project

    return "Not found"


tools = [
    {
        "type": "function",
        "function": {
            "name": "search_resume",
            "description": "Search Kumar Gaurav's parsed resume.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string"
                    }
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_link",
            "description": "Returns portfolio, GitHub, LinkedIn and project links.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string"
                    }
                },
                "required": ["query"]
            }
        }
    }
]

def chat_with_resume():
    print("Welcome Recruiters!")
    print("Type \"exit\" to quit this chat...")
    system_prompt = """
    You are Kumar Gaurav's AI Portfolio Assistant.

    You answer questions only about Kumar Gaurav.

    You have access to two tools:

    1. search_resume(query)
    -- use this whenever the user asks about:
    - skills
    - education
    - experience
    - projects
    - contact information
    - summary
    - certifications
    - technologies
    - achievements

    2. get_link(query)
    -- Use this whenever the user asks for:
    - github
    - linkedin
    - portfolio
    - project links
    - live demo
    - repository

    Rules:
    1. Never invent information.
    2. Use the tools whenever required.
    3. Wait for the tool result before answering.
    4. If a tool returns None or not found, reply:
        "I don't have that information."
    5. Never answer unrelated general knowledge questions.
    6. If greeted, greet politely.
    7. Keep answers concise and professional.
    8. If user ask any other skill that is not in the resume, reply:
        "He doesn't have that skill. but he is always eager to learn new skills and technologies."
    9. If the user asks for a professional summary, profile summary, or "Tell me about Kumar":
    - If a summary exists in the parsed resume, present it.
    - Otherwise, create a professional summary using only the verified resume data.
    - Mention relevant experience, key technical skills, major projects, education, certifications, and career focus when available.
    - Keep the summary concise (4–6 sentences) and professional.
    - Never fabricate information.
    """

    messages = [
        {
            "role": "system",
            "content": system_prompt
        }
    ]

    while True:

        question = input("\nYou: ")

        if question.lower() in ["exit", "quit"]:
            print("Goodbye!")
            break

        messages.append(
            {
                "role": "user",
                "content": question
            }
        )

        response = client.chat.completions.create(
            model=model,
            messages=messages,
            tools=tools,
            temperature=0
        )

        response_message = response.choices[0].message


        # No tool required
        if not response_message.tool_calls:
            answer = response_message.content

            print("\nAssistant: ", answer)

            messages.append(
                {
                    "role": "assistant",
                    "content": answer
                }
            )
            continue

        # tool execution
        messages.append(response_message)

        for tool_call in response_message.tool_calls:
            tool_name = tool_call.function.name

            try:
                arguments = json.loads(tool_call.function.arguments)
            except json.JSONDecodeError:
                arguments = {}

            # execute search_resume()

            if tool_name == "search_resume":
                result = search_resume(
                    arguments["query"]
                )

            # execute get_links()

            elif tool_name == "get_link":
                result = get_link(
                    arguments["query"]
                )

            else:
                result = "Unknown tool."

            messages.append(
                {
                    "role": "tool",
                    "tool_call_id": tool_call.id,
                    "content": json.dumps(
                        result,
                        default=str,
                        indent=2
                    )
                }
            )

        #final llm response

        final_response = client.chat.completions.create(
            model=model,
            messages=messages,
            tools=tools,
            temperature=0
        )

        answer = final_response.choices[0].message.content

        print("\nAssistant: ",answer)

        messages.append(
            {
                "role": "assistant",
                "content": answer
            }
        )


chat_with_resume()